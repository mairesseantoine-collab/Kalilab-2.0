"""
Service d'export DOCX pour les rapports de vérification de méthodes.
- ENR04653 v2 : Méthode quantitative
- ENR04654 v2 : Méthode qualitative
"""
import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL


# ── Couleurs ─────────────────────────────────────────────────────────────────
BLEU_HEADER = RGBColor(0x1F, 0x49, 0x7D)    # bleu foncé titres
GRIS_LABEL = RGBColor(0x40, 0x40, 0x40)     # gris foncé labels
BLANC = RGBColor(0xFF, 0xFF, 0xFF)
BLEU_TABLEAU = RGBColor(0xBD, 0xD7, 0xEE)   # fond en-tête tableau


# ── Helpers XML ───────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_page_number_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS_LABEL


def _add_header(doc: Document, ref: str, titre: str):
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0]
    para.clear()
    run = para.add_run(f"{ref}  –  {titre}")
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS_LABEL
    # ligne séparatrice sous le header
    p_fmt = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    p_fmt.append(pBdr)


# ── Styles de paragraphe ──────────────────────────────────────────────────────
def _h1(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = BLEU_HEADER
    run.font.size = Pt(13)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def _h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = BLEU_HEADER
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def _h3(doc: Document, text: str):
    p = doc.add_heading(text, level=3)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = BLEU_HEADER
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p


def _field(doc: Document, label: str, value: str):
    """Ligne : Label en gras + valeur normale."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r_label = p.add_run(label + " : ")
    r_label.font.bold = True
    r_label.font.size = Pt(10)
    r_val = p.add_run(value or "—")
    r_val.font.size = Pt(10)
    return p


def _block(doc: Document, label: str, value: str):
    """Bloc multiligne : label gras sur ligne seule, puis valeur."""
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.font.bold = True
        r.font.size = Pt(10)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
    if value:
        p2 = doc.add_paragraph(value or "—")
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        for run in p2.runs:
            run.font.size = Pt(10)


def _page_break(doc: Document):
    doc.add_page_break()


# ── Table critères ────────────────────────────────────────────────────────────
def _table_criteres(doc: Document, criteres: List[Dict]):
    headers = ["Paramètre à vérifier", "Critère d'acceptabilité", "Source / Référence"]
    col_widths_cm = [7, 7, 4.5]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # En-tête
    hdr_row = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths_cm)):
        cell = hdr_row.cells[i]
        cell.width = Cm(w)
        _set_cell_bg(cell, "BDD7EE")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLANC
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for c in (criteres or []):
        row = table.add_row()
        vals = [c.get("parametre", ""), c.get("critere", ""), c.get("source", "")]
        for i, (val, w) in enumerate(zip(vals, col_widths_cm)):
            cell = row.cells[i]
            cell.width = Cm(w)
            p = cell.paragraphs[0]
            r = p.add_run(val or "")
            r.font.size = Pt(9)

    doc.add_paragraph()


# ── Section commune Introduction ─────────────────────────────────────────────
def _section_introduction(doc: Document, intro: Dict):
    _h1(doc, "1.\tIntroduction")
    _field(doc, "Nom de la méthode/analyse/équipement", intro.get("nom_methode", ""))
    _field(doc, "Équipe de mise en œuvre", intro.get("equipe", ""))
    _field(doc, "Nom(s) de(s) l'opérateur(s)", intro.get("operateurs", ""))
    ref = intro.get("procedure_ref", "SOPG0003")
    ver = intro.get("version_procedure", "")
    _field(doc, "Procédure de vérification/validation", f"{ref} – version {ver}" if ver else ref)
    debut = intro.get("periode_debut", "")
    fin = intro.get("periode_fin", "")
    _field(doc, "Période d'étude", f"du {debut} au {fin}" if (debut or fin) else "")


def _section_description(doc: Document, desc: Dict):
    _h1(doc, "2.\tDescription de la méthode/analyse/équipement")
    for label, key in [
        ("Nom de la méthode/analyse", "nom_methode"),
        ("Type de méthode", "type"),
        ("Principe de la mesure/Technique", "principe"),
        ("Matrice (+ volume) et stabilité", "matrice"),
        ("Prétraitement de l'échantillon", "pretraitement"),
        ("Unités", "unites"),
        ("Intervalles de référence / Critères d'interprétation", "intervalles_reference"),
        ("Automate", "automate"),
        ("Kit/Matériels et réactifs utilisés", "kit_reactifs"),
        ("Matériau d'étalonnage (références)", "materiau_etalonnage"),
        ("Type d'étalonnage, nombre de niveaux et valeurs", "type_etalonnage"),
        ("Marquage CE-IVDR", "marquage_ce"),
        ("Localisation", "localisation"),
        ("Nombre d'analyses effectuées/an", "nb_analyses_an"),
        ("Objectif de l'analyse / Utilisation prévue", "objectif"),
    ]:
        _field(doc, label, desc.get(key, ""))


def _section_motivation(doc: Document, texte: str):
    _h1(doc, "3.\tMotivation du changement / introduction d'une nouvelle méthode")
    p = doc.add_paragraph(texte or "")
    for run in p.runs:
        run.font.size = Pt(10)


def _section_litterature(doc: Document, texte: str):
    _h1(doc, "4.\tLittérature")
    p = doc.add_paragraph(texte or "")
    for run in p.runs:
        run.font.size = Pt(10)


def _section_criteres(doc: Document, criteres: List[Dict]):
    _h1(doc, "5.\tExigences / Critères d'acceptabilité")
    p = doc.add_paragraph(
        "Les exigences/critères d'acceptabilité ont été décrits dans le tableau ci-dessous "
        "reprenant les paramètres à vérifier et/ou à connaître."
    )
    for run in p.runs:
        run.font.size = Pt(10)
    _table_criteres(doc, criteres)


def _section_resultats_qualitatif(doc: Document, res: Dict):
    _h1(doc, "6.\tRésultats de l'évaluation des performances de la méthode")
    p = doc.add_paragraph(
        "Pour une vérification/validation d'une méthode qualitative (« portée A »), "
        "le dossier portera sur les points suivants :"
    )
    for run in p.runs:
        run.font.size = Pt(10)
    _field(doc, "Données brutes conservées", res.get("donnees_brutes", ""))

    _h2(doc, "6.1\tSpécificité analytique (ou réactions croisées)")
    _block(doc, "", res.get("specificite_analytique", ""))

    _h2(doc, "6.2\tSpécificité diagnostique ou clinique")
    _block(doc, "", res.get("specificite_diagnostique", ""))

    _h2(doc, "6.3\tSensibilité analytique")
    _block(doc, "", res.get("sensibilite_analytique", ""))

    _h2(doc, "6.4\tSensibilité diagnostique ou clinique")
    _block(doc, "", res.get("sensibilite_diagnostique", ""))

    _h2(doc, "6.5\tFidélité de mesure")
    _h3(doc, "6.5.1\tRépétabilité")
    _block(doc, "", res.get("fidelite_repetabilite", ""))
    _h3(doc, "6.5.2\tFidélité intermédiaire")
    _block(doc, "", res.get("fidelite_intermediaire", ""))

    _h2(doc, "6.6\tJustesse")
    _block(doc, "", res.get("justesse", ""))

    _h2(doc, "6.7\tExactitude")
    _block(doc, "", res.get("exactitude", ""))

    _h2(doc, "6.8\tContamination")
    _block(doc, "", res.get("contamination", ""))

    _h2(doc, "6.9\tInterférences")
    _block(doc, "", res.get("interferences", ""))

    _h2(doc, "6.10\tStabilité des réactifs")
    _block(doc, "", res.get("stabilite_reactifs", ""))

    _h2(doc, "6.11\tRobustesse")
    _block(doc, "", res.get("robustesse", ""))

    _h2(doc, "6.12\tComparaison de méthodes")
    _h3(doc, "Méthode précédente, autre méthode utilisée dans le laboratoire :")
    _block(doc, "", res.get("comparaison_methode_precedente", ""))
    _h3(doc, "Nombre de mesures/d'échantillons et descriptif de l'échantillon étudié :")
    _block(doc, "", res.get("comparaison_nb_mesures", ""))
    _h3(doc, "Méthode d'exploitation des résultats (étude des concordances) :")
    _block(doc, "", res.get("comparaison_exploitation", ""))
    _h3(doc, "Résultats et interprétations des discordances :")
    _block(doc, "", res.get("comparaison_discordances", ""))

    _h2(doc, "6.13\tComparaison des résultats d'un même paramètre venant d'autres méthodes de mesure")
    _block(doc, "", res.get("comparaison_autres", ""))

    _h2(doc, "6.14\tIntervalle de référence")
    _block(doc, "", res.get("intervalle_reference", ""))

    _h2(doc, "6.15\tIncertitude de mesure / facteurs de variabilité")
    _block(doc, "", res.get("incertitude", ""))


def _section_resultats_quantitatif(doc: Document, res: Dict):
    _h1(doc, "6.\tRésultats de l'évaluation des performances de la méthode")
    p = doc.add_paragraph(
        "Pour une vérification/validation d'une méthode quantitative (« portée A »), "
        "le dossier portera sur les points suivants :"
    )
    for run in p.runs:
        run.font.size = Pt(10)
    _field(doc, "Données brutes conservées", res.get("donnees_brutes", ""))

    _h2(doc, "6.1\tSpécificité analytique (ou réactions croisées)")
    _block(doc, "", res.get("specificite_analytique", ""))

    _h2(doc, "6.2\tSpécificité diagnostique ou clinique")
    _block(doc, "", res.get("specificite_diagnostique", ""))

    _h2(doc, "6.3\tSensibilité analytique")
    _block(doc, "", res.get("sensibilite_analytique", ""))

    _h2(doc, "6.4\tSensibilité diagnostique ou clinique")
    _block(doc, "", res.get("sensibilite_diagnostique", ""))

    _h2(doc, "6.5\tFidélité de mesure")
    _h3(doc, "6.5.1\tRépétabilité")
    _block(doc, "", res.get("fidelite_repetabilite", ""))
    _h3(doc, "6.5.2\tFidélité intermédiaire")
    _block(doc, "", res.get("fidelite_intermediaire", ""))

    _h2(doc, "6.6\tJustesse")
    _block(doc, "", res.get("justesse", ""))

    _h2(doc, "6.7\tExactitude")
    _block(doc, "", res.get("exactitude", ""))

    _h2(doc, "6.8\tContamination")
    _block(doc, "", res.get("contamination", ""))

    _h2(doc, "6.9\tInterférences")
    _block(doc, "", res.get("interferences", ""))

    _h2(doc, "6.10\tStabilité des réactifs")
    _block(doc, "", res.get("stabilite_reactifs", ""))

    _h2(doc, "6.11\tRobustesse")
    _block(doc, "", res.get("robustesse", ""))

    _h2(doc, "6.12\tComparaison de méthodes")
    _h3(doc, "Méthode précédente, autre méthode utilisée dans le laboratoire :")
    _block(doc, "", res.get("comparaison_methode_precedente", ""))
    _h3(doc, "Nombre de mesures/d'échantillons et descriptif de l'échantillon étudié :")
    _block(doc, "", res.get("comparaison_nb_mesures", ""))
    _h3(doc, "Méthode d'exploitation des résultats (étude des concordances) :")
    _block(doc, "", res.get("comparaison_exploitation", ""))
    _h3(doc, "Résultats et interprétations des discordances :")
    _block(doc, "", res.get("comparaison_discordances", ""))

    _h2(doc, "6.13\tComparaison des résultats d'un même paramètre venant d'autres méthodes de mesure")
    _block(doc, "", res.get("comparaison_autres", ""))

    _h2(doc, "6.14\tIntervalle de référence")
    _block(doc, "", res.get("intervalle_reference", ""))

    _h2(doc, "6.15\tIncertitude de mesure / facteurs de variabilité")
    _block(doc, "", res.get("incertitude", ""))

    # Section exclusive quantitatif
    _h2(doc, "6.16\tIntervalle de mesure")
    _h3(doc, "Limites de détection")
    _block(doc, "", res.get("intervalle_mesure_lod", ""))
    _h3(doc, "Limites de quantification")
    _block(doc, "", res.get("intervalle_mesure_loq", ""))
    _h3(doc, "Limites de linéarité")
    _block(doc, "", res.get("intervalle_mesure_linearite", ""))


def _section_risques(doc: Document, texte: str, num: int):
    _h1(doc, f"{num}.\tMaîtrise / Analyse des risques")
    p = doc.add_paragraph(texte or "")
    for run in p.runs:
        run.font.size = Pt(10)


def _section_validation_informatique(doc: Document, texte: str, num: int):
    _h1(doc, f"{num}.\tValidation informatique et vérification du transfert électronique")
    p = doc.add_paragraph(texte or "")
    for run in p.runs:
        run.font.size = Pt(10)


def _section_conclusion(doc: Document, conc: Dict, num: int):
    _h1(doc, f"{num}.\tConclusion")
    _field(doc, "La méthode répond à nos exigences/critères d'acceptabilité",
           conc.get("repond_exigences", "OUI/NON"))
    _field(doc, "La méthode est validée", conc.get("est_validee", "OUI/NON"))
    _field(doc, "Commentaire(s) éventuel(s)", conc.get("commentaires", ""))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Personne autorisant cette vérification/validation : ")
    run.font.bold = True
    run.font.size = Pt(10)
    doc.add_paragraph().add_run(conc.get("autorisateur", "")).font.size = Pt(10)
    _field(doc, "Date", conc.get("date_autorisation", ""))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(6)
    run2 = p2.add_run("Autorisation par CAQ : ")
    run2.font.bold = True
    run2.font.size = Pt(10)
    doc.add_paragraph().add_run(conc.get("autorisateur_caq", "")).font.size = Pt(10)
    _field(doc, "Date", conc.get("date_caq", ""))


def _section_mise_en_routine(doc: Document, mr: Dict, num: int):
    _h1(doc, f"{num}.\tMise en routine")
    _field(doc, "Date de mise en routine", mr.get("date", ""))
    _field(doc, "Autorisation par", mr.get("autorise_par", ""))


def _section_annexes(doc: Document, texte: str, num: int):
    _h1(doc, f"{num}.\tAnnexes")
    p = doc.add_paragraph(texte or "")
    for run in p.runs:
        run.font.size = Pt(10)


# ── Page de titre ─────────────────────────────────────────────────────────────
def _page_titre(doc: Document, ref: str, type_methode: str, titre: str):
    # Logo / ref document
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_ref.add_run(ref)
    r.font.size = Pt(9)
    r.font.color.rgb = GRIS_LABEL

    # Titre principal
    doc.add_paragraph()
    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titre.paragraph_format.space_before = Pt(60)
    r_t = p_titre.add_run(
        f"Rapport de vérification / validation\nd'une méthode {type_methode}"
    )
    r_t.font.size = Pt(20)
    r_t.font.bold = True
    r_t.font.color.rgb = BLEU_HEADER

    # Sous-titre : nom de la méthode
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(16)
    r_s = p_sub.add_run(titre)
    r_s.font.size = Pt(14)
    r_s.font.color.rgb = GRIS_LABEL

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(24)
    r_d = p_date.add_run(datetime.now().strftime("%d/%m/%Y"))
    r_d.font.size = Pt(11)
    r_d.font.color.rgb = GRIS_LABEL

    _page_break(doc)


# ── Fonctions publiques ───────────────────────────────────────────────────────
def generate_docx_qualitatif(dossier_data: Dict) -> bytes:
    """Génère le rapport ENR04654 (méthode qualitative). Retourne les bytes du .docx."""
    ref = "ENR04654 v2"
    titre = dossier_data.get("titre", "")
    intro = dossier_data.get("introduction", {})
    desc = dossier_data.get("description", {})
    motivation = dossier_data.get("motivation", "")
    litterature = dossier_data.get("litterature", "")
    criteres = dossier_data.get("criteres", [])
    res = dossier_data.get("resultats", {})
    risques = dossier_data.get("risques", "")
    val_info = dossier_data.get("validation_informatique", "")
    conc = dossier_data.get("conclusion", {})
    mr = dossier_data.get("mise_en_routine", {})
    annexes = dossier_data.get("annexes", "")

    doc = Document()
    # Format A4 portrait
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

    # Police par défaut
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    _add_header(doc, ref, titre or "Méthode qualitative")
    _add_page_number_footer(doc)

    _page_titre(doc, ref, "qualitative", titre)

    _section_introduction(doc, intro)
    _section_description(doc, desc)
    _section_motivation(doc, motivation)
    _section_litterature(doc, litterature)
    _section_criteres(doc, criteres)
    _section_resultats_qualitatif(doc, res)
    _section_risques(doc, risques, 7)
    _section_validation_informatique(doc, val_info, 8)
    _section_conclusion(doc, conc, 9)
    _section_mise_en_routine(doc, mr, 10)
    _section_annexes(doc, annexes, 11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_docx_quantitatif(dossier_data: Dict) -> bytes:
    """Génère le rapport ENR04653 (méthode quantitative). Retourne les bytes du .docx."""
    ref = "ENR04653 v2"
    titre = dossier_data.get("titre", "")
    intro = dossier_data.get("introduction", {})
    desc = dossier_data.get("description", {})
    motivation = dossier_data.get("motivation", "")
    litterature = dossier_data.get("litterature", "")
    criteres = dossier_data.get("criteres", [])
    res = dossier_data.get("resultats", {})
    risques = dossier_data.get("risques", "")
    val_info = dossier_data.get("validation_informatique", "")
    conc = dossier_data.get("conclusion", {})
    mr = dossier_data.get("mise_en_routine", {})
    annexes = dossier_data.get("annexes", "")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    _add_header(doc, ref, titre or "Méthode quantitative")
    _add_page_number_footer(doc)

    _page_titre(doc, ref, "quantitative", titre)

    _section_introduction(doc, intro)
    _section_description(doc, desc)
    _section_motivation(doc, motivation)
    _section_litterature(doc, litterature)
    _section_criteres(doc, criteres)
    _section_resultats_quantitatif(doc, res)
    _section_risques(doc, risques, 7)
    _section_validation_informatique(doc, val_info, 8)
    _section_conclusion(doc, conc, 9)
    _section_mise_en_routine(doc, mr, 10)
    _section_annexes(doc, annexes, 11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_docx(dossier_data: Dict) -> bytes:
    """Dispatcher selon type_methode."""
    t = (dossier_data.get("type_methode") or "qualitatif").lower()
    if "quant" in t:
        return generate_docx_quantitatif(dossier_data)
    return generate_docx_qualitatif(dossier_data)
