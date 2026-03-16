"""
Export AMDEC (FMEA) — génération DOCX Python.
Structure calquée sur le modèle Excel du laboratoire (R1-R5).
"""
from io import BytesIO
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Couleurs ──────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1F, 0x49, 0x7D)
BLUE_LIGHT = RGBColor(0xBD, 0xD7, 0xEE)
GREEN      = RGBColor(0xC6, 0xEF, 0xCE)
YELLOW     = RGBColor(0xFF, 0xEB, 0x9C)
RED        = RGBColor(0xFF, 0xC7, 0xCE)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_LIGHT = RGBColor(0xF2, 0xF2, 0xF2)

# IPR thresholds (G × F × D)
IPR_VERT   = 10    # ≤ 10  → vert (acceptable)
IPR_ORANGE = 100   # ≤ 100 → jaune/orange (à surveiller)
# > 100 → rouge (action requise)


def _ipr_color(ipr: int) -> RGBColor:
    if ipr <= IPR_VERT:
        return GREEN
    if ipr <= IPR_ORANGE:
        return YELLOW
    return RED


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, italic=False,
               font_size=7, color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text) if text is not None else "")
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_header_footer(doc: Document, analyseur: str, date_str: str):
    """Entête et pied de page."""
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    htable = header.add_table(1, 3, width=Cm(17))
    htable.style = "Table Grid"
    htable.rows[0].cells[0].text = "ANALYSE DES RISQUES — AMDEC"
    htable.rows[0].cells[1].text = analyseur
    htable.rows[0].cells[2].text = date_str
    for i, cell in enumerate(htable.rows[0].cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8)
                run.font.bold = (i == 0)

    # Footer — numéro de page
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Page ").font.size = Pt(8)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    run.element.append(fldChar1)
    run.element.append(instrText)
    run.element.append(fldChar2)
    run.element.append(fldChar3)


def _info_table(doc: Document, data: Dict):
    """Table d'informations générales (rédacteur, vérificateur, etc.)."""
    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"

    labels = [
        ("Analyseur / Équipement", data.get("analyseur", "")),
        ("Type d'analyseur", data.get("type_analyseur", "")),
        ("Rédacteur", data.get("redacteur", "")),
        ("Date rédaction", data.get("date_redaction", "")),
        ("Vérificateur", data.get("verificateur", "")),
        ("Date vérification", data.get("date_verification", "")),
        ("Validateur", data.get("validateur", "")),
        ("Date validation", data.get("date_validation", "")),
        ("Approbateur (RQ)", data.get("approbateur", "")),
        ("Date approbation", data.get("date_approbation", "")),
        ("Statut", data.get("statut", "brouillon").upper()),
        ("Version", data.get("version", "1.0")),
    ]

    # Repack into 3 rows × 4 cols (2 per cell: label + value)
    rows = tbl.rows
    idx = 0
    for row in rows:
        for j in range(0, 4, 2):
            label_cell = row.cells[j]
            value_cell = row.cells[j + 1] if j + 1 < 4 else row.cells[j]
            if idx < len(labels):
                _cell_text(label_cell, labels[idx][0], bold=True, font_size=8)
                _set_cell_bg(label_cell, GREY_LIGHT)
                _cell_text(value_cell, labels[idx][1], font_size=8)
                idx += 1
    doc.add_paragraph()


def _legend_table(doc: Document):
    """Table légende IPR."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Légende IPR (G × F × D) :")
    run.font.bold = True
    run.font.size = Pt(8)
    run.font.name = "Arial"

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    items = [
        (GREEN,  "IPR ≤ 10 — Zone verte : Risque acceptable, maintien des mesures existantes"),
        (YELLOW, "IPR 11–100 — Zone jaune : Risque à surveiller, action possible"),
        (RED,    "IPR > 100 — Zone rouge : Risque élevé, plan d'action requis"),
    ]
    for i, (color, text) in enumerate(items):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, color)
        _cell_text(cell, text, font_size=7)
    doc.add_paragraph()


CATEGORY_HEADERS = [
    "Sous-item / Risque",
    "Conséquence",
    "Barrières préventives existantes",
    "G", "F", "D", "IPR",
    "Actions à mettre en œuvre",
    "Décision",
    "G'", "F'", "D'", "IPR'",
]

COL_WIDTHS = [
    Cm(3.5),   # Sous-item
    Cm(3.5),   # Conséquence
    Cm(3.5),   # Barrières
    Cm(0.6),   # G
    Cm(0.6),   # F
    Cm(0.6),   # D
    Cm(0.8),   # IPR
    Cm(3.5),   # Actions
    Cm(1.2),   # Décision
    Cm(0.6),   # G'
    Cm(0.6),   # F'
    Cm(0.6),   # D'
    Cm(0.8),   # IPR'
]

DECISION_MAP = {
    "A": "Accepter",
    "R": "Réduire",
    "S": "Supprimer",
    "M": "Maintenir",
    "AS": "À surveiller",
}


def _category_table(doc: Document, category: Dict):
    code = category.get("code", "")
    libelle = category.get("libelle", "")
    items: List[Dict] = category.get("items", [])

    # Category title
    p = doc.add_paragraph()
    run = p.add_run(f"{code} — {libelle}")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BLUE_DARK

    if not items:
        doc.add_paragraph("(Aucun risque enregistré)")
        return

    tbl = doc.add_table(rows=1 + len(items), cols=len(CATEGORY_HEADERS))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    for i, w in enumerate(COL_WIDTHS):
        for row in tbl.rows:
            row.cells[i].width = w

    # Header row
    hrow = tbl.rows[0]
    for i, h in enumerate(CATEGORY_HEADERS):
        _set_cell_bg(hrow.cells[i], BLUE_DARK)
        _cell_text(hrow.cells[i], h, bold=True, font_size=7, color=WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for row_idx, item in enumerate(items):
        G  = int(item.get("G",  1))
        F  = int(item.get("F",  1))
        D  = int(item.get("D",  1))
        Gp = int(item.get("Gp", G))
        Fp = int(item.get("Fp", F))
        Dp = int(item.get("Dp", D))
        ipr     = G  * F  * D
        ipr_p   = Gp * Fp * Dp
        decision_code = item.get("decision", "")
        decision_txt  = DECISION_MAP.get(decision_code, decision_code)

        row = tbl.rows[row_idx + 1]

        values = [
            item.get("sous_item", ""),
            item.get("consequence", ""),
            item.get("barriere", ""),
            str(G), str(F), str(D), str(ipr),
            item.get("action", ""),
            decision_txt,
            str(Gp), str(Fp), str(Dp), str(ipr_p),
        ]

        for col_i, val in enumerate(values):
            cell = row.cells[col_i]
            align = WD_ALIGN_PARAGRAPH.CENTER if col_i in (3, 4, 5, 6, 8, 9, 10, 11, 12) else WD_ALIGN_PARAGRAPH.LEFT
            _cell_text(cell, val, font_size=7, align=align)

        # IPR color (pre-action)
        _set_cell_bg(row.cells[6], _ipr_color(ipr))
        # IPR' color (post-action)
        _set_cell_bg(row.cells[12], _ipr_color(ipr_p))
        # Alternate row bg
        if row_idx % 2 == 1:
            for ci in [0, 1, 2, 7, 8]:
                _set_cell_bg(row.cells[ci], GREY_LIGHT)

    doc.add_paragraph()


def generate_docx_amdec(data: Dict[str, Any]) -> bytes:
    """
    Génère le rapport AMDEC complet en DOCX à partir du dictionnaire de données.
    """
    doc = Document()

    # Page setup A4
    section = doc.sections[0]
    section.page_width  = Cm(29.7)  # A4 landscape
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(1.5)

    analyseur = data.get("analyseur", "Analyseur")
    date_str  = data.get("date_redaction", datetime.today().strftime("%d/%m/%Y"))

    _add_header_footer(doc, analyseur, date_str)

    # ── Titre principal ──
    title = doc.add_heading("ANALYSE DES MODES DE DÉFAILLANCE, DE LEURS EFFETS ET DE LEUR CRITICITÉ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE_DARK

    subtitle = doc.add_paragraph(f"AMDEC — {analyseur}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0] if subtitle.runs else subtitle.add_run(f"AMDEC — {analyseur}")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE_DARK
    doc.add_paragraph()

    # ── Informations générales ──
    h = doc.add_heading("Informations générales", level=2)
    for run in h.runs:
        run.font.color.rgb = BLUE_DARK
        run.font.name = "Arial"
    _info_table(doc, data)

    # ── Légende ──
    _legend_table(doc)

    # ── Catégories de risques ──
    h2 = doc.add_heading("Analyse des risques par catégorie", level=2)
    for run in h2.runs:
        run.font.color.rgb = BLUE_DARK
        run.font.name = "Arial"

    categories = data.get("categories", [])
    for cat in categories:
        _category_table(doc, cat)

    # ── Synthèse ──
    h3 = doc.add_heading("Synthèse et conclusions", level=2)
    for run in h3.runs:
        run.font.color.rgb = BLUE_DARK
        run.font.name = "Arial"

    synthese = data.get("synthese", "")
    p = doc.add_paragraph(synthese if synthese else "À compléter.")
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(9)

    # Return bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
